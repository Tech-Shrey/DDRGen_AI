"""
report_generator.py
-------------------
Generates the final DDR Word document (.docx) from:
  - AI-generated DDR sections
  - Structured property data
  - Extracted inspection + thermal images

Why .docx?
  - Images embed inline with text naturally
  - Client can open/edit in MS Word
  - Easy to convert to PDF with one click
  - python-docx handles formatting cleanly
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


# ── Helpers ───────────────────────────────────────────────────────────────────

def add_heading(doc, text, level=1, color=None):
    """Add a styled heading."""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if color and heading.runs:
        heading.runs[0].font.color.rgb = RGBColor(*color)
    return heading


def add_paragraph(doc, text, bold=False, italic=False, size=11):
    """Add a styled paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return para


def add_image_safe(doc, image_path, width=Inches(3.0), caption=None):
    """
    Add an image to the document safely.
    Skips if file doesn't exist or is too small (likely corrupt).
    """
    if not image_path or not os.path.exists(image_path):
        doc.add_paragraph("[Image Not Available]").italic = True
        return

    # Skip tiny images (under 5KB — likely PDF artifacts not real photos)
    if os.path.getsize(image_path) < 5000:
        return

    try:
        doc.add_picture(image_path, width=width)
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].font.size = Pt(9)
            cap.runs[0].italic = True
    except Exception as e:
        print(f"  [IMAGE ERROR] {os.path.basename(image_path)}: {e}")
        doc.add_paragraph(f"[Image could not be loaded: {os.path.basename(image_path)}]")


def add_divider(doc):
    """Add a horizontal visual divider."""
    doc.add_paragraph("─" * 60)


def clean_ai_text(text):
    """
    Remove markdown bold (**text**) and other markdown
    artifacts from AI output for clean Word formatting.
    """
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)   # bold
    text = re.sub(r'\*(.+?)\*',     r'\1', text)   # italic
    text = re.sub(r'#{1,6}\s+',     '',    text)   # headings
    return text.strip()


# ── Cover Page ────────────────────────────────────────────────────────────────

def build_cover_page(doc, property_info):
    """Build a clean cover page with property details."""

    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style = doc.styles["Title"] if "Title" in [s.name for s in doc.styles] else doc.styles["Normal"]
    run = title.add_run("DETAILED DIAGNOSTIC REPORT")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("Property Inspection & Thermal Analysis")
    run2.font.size = Pt(14)
    run2.italic = True

    doc.add_paragraph()
    add_divider(doc)
    doc.add_paragraph()

    # Property details table
    table = doc.add_table(rows=8, cols=2)
    table.style = "Table Grid"

    details = [
        ("Inspection Date",      property_info.get("inspection_date",  "Not Available")),
        ("Inspected By",         property_info.get("inspected_by",     "Not Available")),
        ("Property Type",        property_info.get("property_type",    "Not Available")),
        ("Floors in Building",   property_info.get("floors",           "Not Available")),
        ("Property Age",         property_info.get("property_age",     "Not Available")),
        ("Inspection Score",     property_info.get("score",            "Not Available")),
        ("Previous Audit Done",  property_info.get("previous_audit",   "Not Available")),
        ("Previous Repair Done", property_info.get("previous_repair",  "Not Available")),
    ]

    for i, (label, value) in enumerate(details):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_page_break()


# ── Section Writers ───────────────────────────────────────────────────────────

def write_section_1(doc, text):
    add_heading(doc, "1. Property Issue Summary", level=1, color=(0x1F, 0x49, 0x7D))
    doc.add_paragraph(clean_ai_text(text))
    doc.add_paragraph()


def write_section_2(doc, text, structured):
    add_heading(doc, "2. Area-Wise Observations", level=1, color=(0x1F, 0x49, 0x7D))
    doc.add_paragraph()

    inspection_images = structured.get("inspection_images", [])
    thermal_map       = structured.get("thermal_map", {})
    impacted_areas    = structured.get("impacted_areas", [])

    # ── Build photo pool: only real photos from appendix pages (11+) ──────────
    import re as _re
    photo_pool = []
    for p in inspection_images:
        size     = os.path.getsize(p)
        basename = os.path.basename(p)
        m        = _re.search(r'page(\d+)', basename)
        if not m:
            continue
        page_num = int(m.group(1))
        if page_num >= 11 and size != 26145 and size > 5000:
            photo_pool.append(p)
    photo_pool = sorted(photo_pool)

    # ── Photo-to-area mapping verified from inspection report source document ──
    # Area 1: Hall           → Photos 1-11  (negative: 1-7,  positive: 8-11)
    # Area 2: Bedroom        → Photos 12-19 (negative: 12-14, positive: 15-19)
    # Area 3: Master Bedroom → Photos 20-30 (negative: 20-25, positive: 26-30)
    # Area 4: Kitchen        → Photos 31-37 (negative: 31-32, positive: 33-37)
    # Area 5: MB Wall        → Photos 38-48 (negative: 38-41, positive: 42-48)
    # Area 6: Parking        → Photos 49-57 (negative: 49-52, positive: 53-57)
    # Area 7: Common Bath    → Photos 58-64 (negative: 58,    positive: 59-64)
    # Index is 0-based into photo_pool (Photo 1 = index 0)
    area_photo_ranges = {
        1: (0,  11),
        2: (11, 19),
        3: (19, 30),
        4: (30, 37),
        5: (37, 48),
        6: (48, 57),
        7: (57, 64),
    }

    # ── Parse AI text into per-area blocks ────────────────────────────────────
    area_texts = {}
    # Split on numbered area patterns like "1. Area 1:" or "Area 1:"
    parts = _re.split(r'(?=(?:\d+\.\s+)?Area\s+\d+[:\s])', clean_ai_text(text))
    for part in parts:
        m = _re.match(r'(?:\d+\.\s+)?Area\s+(\d+)', part.strip())
        if m:
            area_num = int(m.group(1))
            area_texts[area_num] = part.strip()

    # ── Write each area with its photos and thermal scans immediately after ───
    num_areas = max(len(impacted_areas), 7)

    for area_num in range(1, num_areas + 1):

        # Area heading
        add_heading(doc, f"Area {area_num}", level=2, color=(0x2C, 0x5F, 0x2E))

        # Area text from AI
        area_text = area_texts.get(area_num, "")
        if not area_text:
            # Fallback to structured data if AI didn't split cleanly
            area_data = next((a for a in impacted_areas if a["area_number"] == area_num), None)
            if area_data:
                area_text = (
                    f"Problem Observed: {area_data['negative_description']}\n"
                    f"Source / Cause: {area_data['positive_description']}"
                )
        doc.add_paragraph(area_text)
        doc.add_paragraph()

        # ── Inspection photos for this area ───────────────────────────────────
        start, end = area_photo_ranges.get(area_num, (0, 0))
        area_photos = photo_pool[start:end] if start < len(photo_pool) else []

        if area_photos:
            add_heading(doc, "Inspection Photographs", level=3)
            shown = 0
            for img_path in area_photos[:6]:  # max 6 photos per area
                full_path = os.path.abspath(img_path)
                add_image_safe(doc, full_path, width=Inches(2.8),
                               caption=f"Photo {start + shown + 1}")
                shown += 1
                if shown % 2 == 0:
                    doc.add_paragraph()
            doc.add_paragraph()

        # ── Thermal scans for this area ───────────────────────────────────────
        scans = thermal_map.get(area_num, [])
        if scans:
            add_heading(doc, "Thermal Scan Data", level=3)
            for scan in scans[:2]:
                img_path  = scan.get("image_path")
                full_path = os.path.abspath(img_path) if img_path else None
                hotspot   = scan.get("hotspot",    "N/A")
                coldspot  = scan.get("coldspot",   "N/A")
                imgfile   = scan.get("image_file", f"Scan {scan['page']}")
                caption   = f"{imgfile} | Hotspot: {hotspot} | Coldspot: {coldspot}"
                add_image_safe(doc, full_path, width=Inches(3.2), caption=caption)
            doc.add_paragraph()

        doc.add_paragraph()


def write_section_3(doc, text):
    add_heading(doc, "3. Probable Root Cause", level=1, color=(0x1F, 0x49, 0x7D))
    doc.add_paragraph(clean_ai_text(text))
    doc.add_paragraph()


def write_section_4(doc, text):
    add_heading(doc, "4. Severity Assessment", level=1, color=(0x1F, 0x49, 0x7D))
    doc.add_paragraph(clean_ai_text(text))
    doc.add_paragraph()


def write_section_5(doc, text):
    add_heading(doc, "5. Recommended Actions", level=1, color=(0x1F, 0x49, 0x7D))
    doc.add_paragraph(clean_ai_text(text))
    doc.add_paragraph()


def write_section_6(doc, text):
    add_heading(doc, "6. Additional Notes", level=1, color=(0x1F, 0x49, 0x7D))
    doc.add_paragraph(clean_ai_text(text) if text else "None")
    doc.add_paragraph()


def write_section_7(doc, text):
    add_heading(doc, "7. Missing or Unclear Information", level=1, color=(0x1F, 0x49, 0x7D))
    doc.add_paragraph(clean_ai_text(text) if text else "None")
    doc.add_paragraph()


# ── Master Report Builder ─────────────────────────────────────────────────────

def generate_report(sections: dict, structured: dict, output_path: str = "DDR_Report.docx"):
    """
    Build the complete Word document from DDR sections + structured data.

    sections   : output from ai_generator.extract_ddr_sections()
    structured : output from structurer.structure_all()
    output_path: where to save the .docx file
    """
    print("\n=== Generating Word Report ===\n")

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Cover Page ────────────────────────────────────────────────────────────
    build_cover_page(doc, structured["property_info"])

    # ── DDR Sections ─────────────────────────────────────────────────────────
    write_section_1(doc, sections.get("property_issue_summary", "Not Available"))
    write_section_2(doc, sections.get("area_wise_observations", "Not Available"), structured)
    write_section_3(doc, sections.get("probable_root_cause",    "Not Available"))
    write_section_4(doc, sections.get("severity_assessment",    "Not Available"))
    write_section_5(doc, sections.get("recommended_actions",    "Not Available"))
    write_section_6(doc, sections.get("additional_notes",       ""))
    write_section_7(doc, sections.get("missing_unclear_information", ""))

    # ── Save ──────────────────────────────────────────────────────────────────
    doc.save(output_path)
    print(f"Report saved: {output_path}")
    print(f"File size   : {os.path.getsize(output_path) / 1024:.1f} KB")

    return output_path


# ── Quick test ────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    from extractor import extract_all
    from structurer import structure_all
    from prompt_builder import build_prompt
    from ai_generator import generate_ddr, extract_ddr_sections

    # Full pipeline
    extracted  = extract_all("Sample Report.pdf", "Thermal Images.pdf")
    structured = structure_all(extracted)
    prompt     = build_prompt(structured)
    ddr_text   = generate_ddr(prompt)
    sections   = extract_ddr_sections(ddr_text)

    # Generate Word report
    output = generate_report(sections, structured)
    print(f"\nDone! Open '{output}' to view the report.")
